#!/usr/bin/env python3
"""
POC05_v02 - Process 127 Documents with Optimized Chunking
Base: POC04 processing approach (semantic chunking + bilingual metadata)
Output: poc05_v02_optimized_chunks table
"""

import sys
import os
import json
from datetime import datetime

sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'helper_function'))

from databricks.connect import DatabricksSession
from pyspark.sql import Row


def install_libraries():
    """Install required libraries"""
    print("Installing libraries...")
    import subprocess
    libs = ["python-docx", "pythainlp"]
    for lib in libs:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", lib, "--quiet"])
        except:
            pass
    print("Libraries installed")


def extract_keywords(text, lang='en'):
    """Extract keywords from text"""
    words = text.lower().split()

    tech_keywords = {'table', 'column', 'log', 'id', 'status', 'data', 'submission',
                     'response', 'database', 'key', 'index', 'tbl_', 'bot', 'rdt'}

    thai_keywords = {'ตาราง', 'คอลัมน์', 'ข้อมูล', 'สถานะ', 'บันทึก', 'ธปท',
                     'ส่ง', 'ตอบกลับ', 'ระบบ', 'ฐานข้อมูล'}

    found = set()
    for word in words:
        if any(kw in word for kw in tech_keywords):
            found.add(word)
        if any(kw in word for kw in thai_keywords):
            found.add(word)

    return list(found)[:10]


def enhance_chunk_bilingual(chunk_text, section, doc_name, chunk_type='text'):
    """Add bilingual metadata to chunks"""
    en_keywords = extract_keywords(chunk_text, 'en')
    th_keywords = extract_keywords(chunk_text, 'th')

    enhanced = f"""[SECTION] {section}
KEYWORDS: {', '.join(en_keywords)}
คำสำคัญ: {', '.join(th_keywords)}
DOCUMENT: {doc_name}

{chunk_text}"""

    return enhanced


def extract_with_overlap(doc_bytes, filename):
    """
    Enhanced extraction with semantic chunking and overlap
    - Smaller chunks (800 chars max)
    - Overlap (200 chars)
    - Better section detection
    - Bilingual metadata
    """
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(doc_bytes))
        chunks = []
        current_section = "Introduction"
        current_text = []
        current_size = 0
        table_count = 0

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect section headers
            if text.startswith('#') or (para.style.name.startswith('Heading') and len(text) < 100):
                # Save previous section
                if current_text:
                    chunk_content = '\n'.join(current_text)
                    enhanced = enhance_chunk_bilingual(chunk_content, current_section, filename, 'text')
                    chunks.append({
                        'content': enhanced,
                        'content_type': 'text',
                        'section': current_section,
                        'char_count': len(enhanced),
                        'original_char_count': len(chunk_content)
                    })

                # Start new section
                current_section = text.replace('#', '').strip()
                current_text = []
                current_size = 0
                continue

            # Add text to current chunk
            current_text.append(text)
            current_size += len(text)

            # Create chunk if size exceeded
            if current_size > 800:
                chunk_content = '\n'.join(current_text)
                enhanced = enhance_chunk_bilingual(chunk_content, current_section, filename, 'text')
                chunks.append({
                    'content': enhanced,
                    'content_type': 'text',
                    'section': current_section,
                    'char_count': len(enhanced),
                    'original_char_count': len(chunk_content)
                })

                # OVERLAP: Keep last 2 sentences (~200 chars)
                overlap_sentences = current_text[-2:] if len(current_text) > 2 else current_text[-1:]
                current_text = overlap_sentences
                current_size = sum(len(s) for s in current_text)

        # Final text chunk
        if current_text:
            chunk_content = '\n'.join(current_text)
            enhanced = enhance_chunk_bilingual(chunk_content, current_section, filename, 'text')
            chunks.append({
                'content': enhanced,
                'content_type': 'text',
                'section': current_section,
                'char_count': len(enhanced),
                'original_char_count': len(chunk_content)
            })

        # Process tables
        for table in doc.tables:
            table_count += 1
            try:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                table_data = []

                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(' | '.join(row_data))

                if headers and table_data:
                    column_list = ', '.join(headers)
                    table_text = f"""TABLE {table_count}: {current_section}
COLUMNS: {column_list}
Headers: {' | '.join(headers)}
{chr(10).join(table_data[:20])}"""

                    enhanced = enhance_chunk_bilingual(table_text, current_section, filename, 'table')

                    chunks.append({
                        'content': enhanced,
                        'content_type': 'table',
                        'section': current_section,
                        'char_count': len(enhanced),
                        'original_char_count': len(table_text),
                        'table_columns': column_list
                    })

            except Exception as e:
                print(f"      Table {table_count} extraction error: {e}")
                continue

        return chunks

    except Exception as e:
        print(f"    ERROR processing {filename}: {e}")
        return []


def main():
    print("=" * 80)
    print("POC05_V02 - PROCESS 127 DOCUMENTS")
    print("=" * 80)

    install_libraries()

    # Load config
    config_path = os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'databricks_config', 'config.json')
    with open(config_path, 'r') as f:
        config = json.load(f)

    print("\n[1] Connecting to Databricks...")
    spark = DatabricksSession.builder.remote(
        host=config['databricks']['host'],
        token=config['databricks']['token'],
        cluster_id=config['databricks']['cluster_id']
    ).getOrCreate()
    print("    Connected")

    volume_path = config['rag_config']['volume_path']
    print(f"\n[2] Scanning documents in {volume_path}...")

    # List documents
    files_df = spark.sql(f"LIST '{volume_path}'")
    all_files = [row['name'] for row in files_df.collect()]

    # Filter document files
    doc_files = [f for f in all_files if f.lower().endswith(('.docx', '.pdf', '.doc'))]
    print(f"    Found {len(doc_files)} documents (docx/pdf/doc only)")

    print(f"\n[3] Processing documents with optimizations...")
    print(f"    - Semantic chunking (800 char max)")
    print(f"    - 200 char overlap for context")
    print(f"    - Bilingual keywords (Thai + English)")
    print(f"    - Table column extraction")

    all_chunks = []
    processed_count = 0
    text_chunks_list = []
    table_chunks_list = []

    for filename in doc_files:
        # Skip non-docx for now (PDF needs different handling)
        if not filename.lower().endswith('.docx'):
            continue

        try:
            file_path = f"{volume_path}/{filename}"

            # Read file
            file_df = spark.read.format("binaryFile").load(file_path)
            file_bytes = file_df.collect()[0]['content']

            # Process
            chunks = extract_with_overlap(file_bytes, filename)

            if chunks:
                for i, chunk in enumerate(chunks):
                    chunk_id = f"poc05_v02_{processed_count}_{i}"
                    chunk_row = {
                        'id': chunk_id,
                        'content': chunk['content'],
                        'content_type': chunk['content_type'],
                        'section': chunk['section'],
                        'char_count': chunk['char_count'],
                        'original_char_count': chunk['original_char_count'],
                        'source_document': filename,
                        'table_columns': chunk.get('table_columns', None),
                        'processing_method': 'semantic_chunking_overlap_bilingual'
                    }
                    all_chunks.append(chunk_row)

                    if chunk['content_type'] == 'text':
                        text_chunks_list.append(chunk)
                    else:
                        table_chunks_list.append(chunk)

                processed_count += 1
                print(f"  Processing: {filename}")
                print(f"    Extracted {len(chunks)} chunks (optimized with overlap & bilingual metadata)")

        except Exception as e:
            print(f"  Processing: {filename}")
            print(f"    ERROR processing {filename}: {e}")
            continue

    # Save to table
    print(f"\n[4] Processing complete:")
    print(f"    Documents processed: {processed_count}")
    print(f"    Total chunks: {len(all_chunks)}")

    if text_chunks_list:
        avg_text_chars = sum(c['original_char_count'] for c in text_chunks_list) // len(text_chunks_list)
        print(f"    Text chunks: {len(text_chunks_list)} (avg {avg_text_chars} chars)")

    if table_chunks_list:
        avg_table_chars = sum(c['original_char_count'] for c in table_chunks_list) // len(table_chunks_list)
        print(f"    Table chunks: {len(table_chunks_list)} (avg {avg_table_chars} chars)")

    # Create DataFrame
    print(f"\n[5] Saving to table: sandbox.rdt_knowledge.poc05_v02_optimized_chunks")
    chunks_df = spark.createDataFrame([Row(**chunk) for chunk in all_chunks])
    chunks_df.write.format("delta").mode("overwrite").saveAsTable("sandbox.rdt_knowledge.poc05_v02_optimized_chunks")
    print(f"    Saved {len(all_chunks)} chunks")

    # Update config
    config['rag_config']['poc05_v02_table'] = 'sandbox.rdt_knowledge.poc05_v02_optimized_chunks'
    config['rag_config']['poc05_v02_chunk_count'] = len(all_chunks)
    config['rag_config']['poc05_v02_created_at'] = datetime.now().isoformat()

    with open(config_path, 'w') as f:
        json.dump(config, f, indent=2)

    print("\n" + "=" * 80)
    print("POC05_V02 PROCESSING COMPLETE")
    print("=" * 80)
    print(f"\nTable: sandbox.rdt_knowledge.poc05_v02_optimized_chunks")
    print(f"Chunks: {len(all_chunks)}")
    print(f"Documents: {processed_count}/{len(doc_files)}")
    print(f"\nNext: Add metadata columns with 03b_add_metadata_sql.py")

    spark.stop()


if __name__ == "__main__":
    main()

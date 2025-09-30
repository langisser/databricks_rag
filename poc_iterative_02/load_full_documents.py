#!/usr/bin/env python3
"""
Load FULL document content from Volume files
- Direct file access from Unity Catalog volume
- Extract actual text from DOCX and Excel files
- Process complete documents, not samples
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'helper_function'))

from databricks.connect import DatabricksSession
import json
import time
import tempfile
import requests

def download_file_from_volume(spark, volume_path, local_temp_path):
    """Download file from volume to local temp for processing"""
    try:
        # Read file as binary from volume
        binary_df = spark.read.format("binaryFile").load(volume_path)
        file_data = binary_df.collect()[0]

        # Write to temporary local file
        with open(local_temp_path, 'wb') as f:
            f.write(file_data['content'])

        return True
    except Exception as e:
        print(f"    Download failed: {e}")
        return False

def extract_docx_full_content(local_file_path):
    """Extract full text from DOCX file"""
    try:
        from docx import Document

        doc = Document(local_file_path)
        full_text = []

        # Extract all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())

        content = "\\n".join(full_text)
        return content, "docx_full_extraction"

    except Exception as e:
        return f"[DOCX EXTRACTION ERROR: {str(e)}]", "docx_failed"

def extract_excel_full_content(local_file_path):
    """Extract full content from Excel file"""
    try:
        import pandas as pd

        # Read all sheets
        excel_file = pd.ExcelFile(local_file_path)
        all_content = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(local_file_path, sheet_name=sheet_name)

            # Add sheet header
            all_content.append(f"Sheet: {sheet_name}")

            # Add column headers
            headers = " | ".join([str(col) for col in df.columns])
            all_content.append(f"Columns: {headers}")

            # Add data rows (convert to string and handle NaN)
            for index, row in df.iterrows():
                row_data = []
                for value in row:
                    if pd.isna(value):
                        row_data.append("")
                    else:
                        row_data.append(str(value))

                if any(cell.strip() for cell in row_data):  # Only add non-empty rows
                    all_content.append(" | ".join(row_data))

            all_content.append("")  # Empty line between sheets

        content = "\\n".join(all_content)
        return content, "excel_full_extraction"

    except Exception as e:
        return f"[EXCEL EXTRACTION ERROR: {str(e)}]", "excel_failed"

def main():
    print("=" * 60)
    print("LOADING FULL DOCUMENT CONTENT FROM VOLUME")
    print("=" * 60)

    try:
        # Load configuration
        config_path = os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'databricks_config', 'config.json')
        with open(config_path, 'r') as f:
            config = json.load(f)

        spark = DatabricksSession.builder.remote(
            host=config['databricks']['host'],
            token=config['databricks']['token'],
            cluster_id=config['databricks']['cluster_id']
        ).getOrCreate()

        namespace = config['rag_config']['full_namespace']
        volume_path = "/Volumes/sandbox/rdt_knowledge/rdt_document"

        print(f"Volume: {volume_path}")
        print(f"Namespace: {namespace}")

        # Get your files
        files_result = spark.sql(f"LIST '{volume_path}'")
        files = files_result.collect()

        print(f"\\nFound {len(files)} files to process with FULL content extraction")

        # Drop and recreate tables for full content
        print(f"\\nRecreating tables for full content...")

        spark.sql(f"DROP TABLE IF EXISTS {namespace}.full_documents")
        spark.sql(f"DROP TABLE IF EXISTS {namespace}.full_document_chunks")

        # Create tables for full content
        spark.sql(f"""
            CREATE TABLE {namespace}.full_documents (
                document_id STRING,
                file_name STRING,
                file_path STRING,
                file_type STRING,
                file_size_kb DOUBLE,
                full_content STRING,
                content_length INT,
                word_count INT,
                paragraph_count INT,
                extraction_method STRING,
                processing_status STRING,
                created_timestamp TIMESTAMP,
                metadata MAP<STRING, STRING>
            )
            USING DELTA
        """)

        spark.sql(f"""
            CREATE TABLE {namespace}.full_document_chunks (
                chunk_id STRING,
                document_id STRING,
                chunk_index INT,
                chunk_text STRING,
                chunk_word_count INT,
                chunk_char_count INT,
                chunk_strategy STRING,
                semantic_density DOUBLE,
                created_timestamp TIMESTAMP,
                metadata MAP<STRING, STRING>
            )
            USING DELTA
        """)

        print("Tables created for full content")

        # Process each file with FULL content extraction
        processed_count = 0
        total_chunks = 0

        for file_info in files:
            file_name = file_info['name']
            file_path = file_info['path']
            file_size = file_info['size']
            file_type = file_name.split('.')[-1].lower() if '.' in file_name else 'unknown'

            print(f"\\nProcessing FULL content: {file_name}")
            print(f"  Type: {file_type}, Size: {round(file_size/1024, 1)} KB")

            doc_id = f"full_doc_{int(time.time())}_{processed_count + 1}"

            # Create temporary directory for file processing
            with tempfile.TemporaryDirectory() as temp_dir:
                local_file_path = os.path.join(temp_dir, file_name)

                # Download file from volume to local temp
                print(f"  Downloading from volume...")
                if not download_file_from_volume(spark, file_path, local_file_path):
                    print(f"  SKIP: Could not download {file_name}")
                    continue

                # Extract full content based on file type
                content = ""
                extraction_method = ""

                if file_type == 'docx':
                    print(f"  Extracting FULL DOCX content...")
                    content, extraction_method = extract_docx_full_content(local_file_path)

                elif file_type == 'xlsx':
                    print(f"  Extracting FULL Excel content...")
                    content, extraction_method = extract_excel_full_content(local_file_path)

                else:
                    # Try to read as text file
                    try:
                        with open(local_file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        extraction_method = "text_file_read"
                    except:
                        try:
                            with open(local_file_path, 'r', encoding='latin-1') as f:
                                content = f.read()
                            extraction_method = "text_file_latin1"
                        except Exception as e:
                            content = f"[COULD NOT EXTRACT: {str(e)}]"
                            extraction_method = "extraction_failed"

                # Calculate content metrics
                if content and not content.startswith('['):  # Successful extraction
                    word_count = len(content.split())
                    char_count = len(content)
                    paragraph_count = len([p for p in content.split('\\n') if p.strip()])
                    processing_status = "success"

                    print(f"  SUCCESS: Extracted {word_count} words, {char_count} characters")
                    print(f"  Paragraphs: {paragraph_count}")
                else:
                    word_count = 0
                    char_count = len(content) if content else 0
                    paragraph_count = 0
                    processing_status = "failed"
                    print(f"  FAILED: {content[:100]}...")

                # Insert full document
                spark.sql(f"""
                    INSERT INTO {namespace}.full_documents VALUES (
                        '{doc_id}',
                        '{file_name}',
                        '{file_path}',
                        '{file_type}',
                        {file_size / 1024:.2f},
                        '{content.replace("'", "''") if content else ""}',
                        {char_count},
                        {word_count},
                        {paragraph_count},
                        '{extraction_method}',
                        '{processing_status}',
                        current_timestamp(),
                        map(
                            'original_size_bytes', '{file_size}',
                            'temp_file_used', 'true',
                            'full_extraction', 'true'
                        )
                    )
                """)

                # Create chunks from full content
                if processing_status == "success" and word_count > 0:
                    print(f"  Creating chunks from full content...")

                    # Determine chunk strategy based on content size
                    if word_count < 500:
                        chunk_size, overlap = 300, 50
                    elif word_count < 2000:
                        chunk_size, overlap = 500, 100
                    elif word_count < 5000:
                        chunk_size, overlap = 800, 150
                    else:
                        chunk_size, overlap = 1000, 200

                    # Create semantic chunks
                    import re

                    # Clean content
                    clean_content = re.sub(r'\\n+', ' ', content)
                    clean_content = re.sub(r'\\s+', ' ', clean_content).strip()

                    # Split into sentences
                    sentences = re.split(r'[.!?]+\\s+', clean_content)
                    sentences = [s.strip() for s in sentences if s.strip()]

                    # Build chunks
                    chunks = []
                    current_chunk = ""
                    current_words = 0
                    chunk_index = 0

                    for sentence in sentences:
                        sentence_words = len(sentence.split())

                        if current_words + sentence_words > chunk_size and current_chunk:
                            chunks.append({
                                "text": current_chunk.strip(),
                                "word_count": current_words,
                                "index": chunk_index
                            })

                            # Start new chunk with overlap
                            if overlap > 0 and current_words > overlap:
                                overlap_words = current_chunk.split()[-overlap:]
                                current_chunk = " ".join(overlap_words) + " " + sentence
                                current_words = len(overlap_words) + sentence_words
                            else:
                                current_chunk = sentence
                                current_words = sentence_words

                            chunk_index += 1
                        else:
                            if current_chunk:
                                current_chunk += " " + sentence
                            else:
                                current_chunk = sentence
                            current_words += sentence_words

                    # Add final chunk
                    if current_chunk:
                        chunks.append({
                            "text": current_chunk.strip(),
                            "word_count": current_words,
                            "index": chunk_index
                        })

                    # Insert chunks
                    for chunk in chunks:
                        chunk_id = f"{doc_id}_chunk_{chunk['index']:03d}"
                        chunk_text = chunk['text']
                        chunk_words = chunk['word_count']
                        chunk_chars = len(chunk_text)

                        # Calculate semantic density
                        words = chunk_text.lower().split()
                        unique_words = len(set(words))
                        semantic_density = unique_words / len(words) if words else 0

                        spark.sql(f"""
                            INSERT INTO {namespace}.full_document_chunks VALUES (
                                '{chunk_id}',
                                '{doc_id}',
                                {chunk['index']},
                                '{chunk_text.replace("'", "''")}',
                                {chunk_words},
                                {chunk_chars},
                                '{chunk_size}w_{overlap}o',
                                {semantic_density:.4f},
                                current_timestamp(),
                                map(
                                    'source_file', '{file_name}',
                                    'extraction_method', '{extraction_method}',
                                    'full_content', 'true'
                                )
                            )
                        """)

                    total_chunks += len(chunks)
                    print(f"  Created {len(chunks)} chunks from full content")

                processed_count += 1

        # Show results
        print(f"\\n" + "=" * 60)
        print("FULL CONTENT EXTRACTION RESULTS")
        print("=" * 60)

        # Document stats
        doc_stats = spark.sql(f"""
            SELECT
                file_type,
                COUNT(*) as doc_count,
                AVG(word_count) as avg_words,
                SUM(word_count) as total_words,
                AVG(content_length) as avg_chars
            FROM {namespace}.full_documents
            WHERE processing_status = 'success'
            GROUP BY file_type
        """).collect()

        print("Full Document Statistics:")
        for stat in doc_stats:
            print(f"  {stat['file_type']}: {stat['doc_count']} docs, {stat['total_words']:.0f} total words")
            print(f"    Average: {stat['avg_words']:.0f} words, {stat['avg_chars']:.0f} characters")

        # Chunk stats
        chunk_total = spark.sql(f"SELECT COUNT(*) as count FROM {namespace}.full_document_chunks").collect()[0]['count']

        print(f"\\nFull Content Processing Summary:")
        print(f"  Documents processed: {processed_count}")
        print(f"  Total chunks created: {chunk_total}")
        print(f"  Tables: full_documents, full_document_chunks")

        spark.stop()

        print(f"\\nSUCCESS: Your documents are now loaded with FULL content!")
        print(f"Ready for vector search indexing with complete document content.")

        return True

    except Exception as e:
        print(f"\\nERROR: Full content extraction failed: {e}")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
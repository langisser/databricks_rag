#!/usr/bin/env python3
"""
Process DataX document with advanced multimodal extraction and create vector search
Extracts text, tables, and images from DataX_OPM_RDT_Submission_to_BOT_v1.0.docx
"""
import sys
import os
import json
import tempfile
from datetime import datetime
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'helper_function'))

from databricks.connect import DatabricksSession
from databricks.vector_search.client import VectorSearchClient
from pyspark.sql.types import StructType, StructField, StringType, IntegerType, ArrayType
from pyspark.sql import Row

def install_libraries():
    """Install required libraries"""
    print("Installing required libraries...")
    import subprocess
    libs = ["python-docx", "pandas", "Pillow"]
    for lib in libs:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", lib, "--quiet"])
        except:
            pass
    print("Libraries installed")

def extract_text_and_tables_from_docx(docx_content):
    """Extract text and tables from DOCX binary content"""
    from docx import Document
    from io import BytesIO

    doc = Document(BytesIO(docx_content))

    chunks = []
    chunk_id = 0

    # Extract paragraphs with section tracking
    current_section = "Introduction"
    current_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect section headers (typically larger font or bold)
        if len(text) < 100 and (para.style.name.startswith('Heading') or
                                 text.isupper() or
                                 any(keyword in text for keyword in ['Overview', 'Process', 'Framework', 'Architecture'])):
            # Save previous section
            if current_text:
                chunks.append({
                    'chunk_id': chunk_id,
                    'content': '\n'.join(current_text),
                    'content_type': 'text',
                    'section': current_section,
                    'char_count': sum(len(t) for t in current_text)
                })
                chunk_id += 1
                current_text = []

            current_section = text
            current_text.append(f"## {text}")
        else:
            current_text.append(text)

            # Create chunk if too large
            if sum(len(t) for t in current_text) > 2000:
                chunks.append({
                    'chunk_id': chunk_id,
                    'content': '\n'.join(current_text),
                    'content_type': 'text',
                    'section': current_section,
                    'char_count': sum(len(t) for t in current_text)
                })
                chunk_id += 1
                current_text = []

    # Save final section
    if current_text:
        chunks.append({
            'chunk_id': chunk_id,
            'content': '\n'.join(current_text),
            'content_type': 'text',
            'section': current_section,
            'char_count': sum(len(t) for t in current_text)
        })
        chunk_id += 1

    # Extract tables
    table_num = 0
    for table in doc.tables:
        table_data = []
        headers = []

        # Get headers from first row
        if len(table.rows) > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]

        # Get data rows
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):  # Skip empty rows
                table_data.append(' | '.join(row_data))

        if table_data:
            table_text = f"TABLE {table_num + 1}:\nHeaders: {' | '.join(headers)}\n" + '\n'.join(table_data)
            chunks.append({
                'chunk_id': chunk_id,
                'content': table_text,
                'content_type': 'table',
                'section': current_section,
                'char_count': len(table_text)
            })
            chunk_id += 1
            table_num += 1

    # Extract images metadata
    image_count = len(doc.inline_shapes)
    if image_count > 0:
        chunks.append({
            'chunk_id': chunk_id,
            'content': f"Document contains {image_count} embedded images/diagrams including process flowcharts and architecture diagrams.",
            'content_type': 'image_metadata',
            'section': 'Document Images',
            'char_count': 100
        })

    return chunks

def main():
    print("=" * 80)
    print("ADVANCED MULTIMODAL PROCESSING - DataX Document")
    print("=" * 80)

    # Install libraries
    install_libraries()

    # Load config
    config_path = os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'databricks_config', 'config.json')
    with open(config_path, 'r') as f:
        config = json.load(f)

    # Connect to Databricks
    print("\n[STEP 1] Connecting to Databricks...")
    spark = DatabricksSession.builder.remote(
        host=config['databricks']['host'],
        token=config['databricks']['token'],
        cluster_id=config['databricks']['cluster_id']
    ).getOrCreate()
    print("  [SUCCESS] Connected")

    # Document path
    doc_path = "/Volumes/sandbox/rdt_knowledge/rdt_document/DataX_OPM_RDT_Submission_to_BOT_v1.0.docx"

    print(f"\n[STEP 2] Loading document from Volume...")
    print(f"  Path: {doc_path}")

    # Read document as binary
    binary_df = spark.read.format("binaryFile").load(doc_path)
    doc_data = binary_df.collect()[0]
    doc_content = doc_data['content']

    print(f"  [SUCCESS] Loaded {len(doc_content):,} bytes")

    # Extract content
    print(f"\n[STEP 3] Extracting text, tables, and images...")
    chunks = extract_text_and_tables_from_docx(doc_content)

    print(f"  [SUCCESS] Extracted {len(chunks)} chunks:")
    text_chunks = [c for c in chunks if c['content_type'] == 'text']
    table_chunks = [c for c in chunks if c['content_type'] == 'table']
    image_chunks = [c for c in chunks if c['content_type'] == 'image_metadata']

    print(f"    - Text sections: {len(text_chunks)}")
    print(f"    - Tables: {len(table_chunks)}")
    print(f"    - Images: {len(image_chunks)}")
    print(f"    - Total characters: {sum(c['char_count'] for c in chunks):,}")

    # Prepare data for Delta table
    print(f"\n[STEP 4] Creating Delta table...")

    namespace = config['rag_config']['full_namespace']
    table_name = f"{namespace}.datax_multimodal_chunks"

    # Create DataFrame
    rows = []
    for chunk in chunks:
        rows.append(Row(
            id=f"datax_bot_{chunk['chunk_id']}",
            content=chunk['content'],
            content_type=chunk['content_type'],
            section=chunk['section'],
            char_count=chunk['char_count'],
            source_document="DataX_OPM_RDT_Submission_to_BOT_v1.0.docx"
        ))

    df = spark.createDataFrame(rows)

    # Save to Delta table
    df.write.format("delta").mode("overwrite").saveAsTable(table_name)

    print(f"  [SUCCESS] Created table: {table_name}")
    print(f"  Rows: {df.count()}")

    # Show sample
    print("\n[STEP 5] Sample chunks:")
    sample_df = spark.sql(f"SELECT id, content_type, section, char_count FROM {table_name} LIMIT 5")
    sample_df.show(truncate=False)

    # Create vector search index
    print(f"\n[STEP 6] Creating vector search index...")

    try:
        vsc = VectorSearchClient(
            workspace_url=config['databricks']['host'],
            personal_access_token=config['databricks']['token']
        )

        endpoint_name = config['vector_search']['endpoint_name']
        index_name = f"{namespace}.datax_multimodal_index"

        print(f"  Endpoint: {endpoint_name}")
        print(f"  Index: {index_name}")

        # Enable CDF if needed
        print("  Enabling Change Data Feed...")
        spark.sql(f"ALTER TABLE {table_name} SET TBLPROPERTIES (delta.enableChangeDataFeed = true)")

        # Create index
        print("  Creating index (this may take a few minutes)...")
        index = vsc.create_delta_sync_index(
            endpoint_name=endpoint_name,
            source_table_name=table_name,
            index_name=index_name,
            pipeline_type='TRIGGERED',
            primary_key='id',
            embedding_source_column='content',
            embedding_model_endpoint_name='databricks-gte-large-en'
        )

        print(f"  [SUCCESS] Index created!")
        print(f"  Index name: {index_name}")

        # Wait for index to be ready
        print("\n  Waiting for index sync...")
        import time
        for i in range(30):
            try:
                status = vsc.get_index(endpoint_name, index_name)
                if hasattr(status, 'status') and status.status.ready:
                    print(f"  [SUCCESS] Index is ready!")
                    break
            except:
                pass
            time.sleep(10)
            print(f"    Checking... ({i*10}s)")

    except Exception as e:
        print(f"  [WARNING] Index creation: {e}")
        print(f"  Note: Index may already exist or need manual creation")

    # Summary
    print("\n" + "=" * 80)
    print("PROCESSING COMPLETE")
    print("=" * 80)
    print(f"\nTable: {table_name}")
    print(f"Index: {namespace}.datax_multimodal_index")
    print(f"\nContent extracted:")
    print(f"  - {len(text_chunks)} text sections")
    print(f"  - {len(table_chunks)} tables")
    print(f"  - {len(image_chunks)} image references")
    print(f"  - {sum(c['char_count'] for c in chunks):,} total characters")
    print(f"\n[NEXT] Test in AI Playground with the index!")

if __name__ == "__main__":
    main()
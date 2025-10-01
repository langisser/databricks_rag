#!/usr/bin/env python3
"""
POC05_v02 - Process ALL Document Formats (docx, pdf, xlsx)
Extends 02_process script to handle:
- .docx files (Word documents)
- .pdf files (PDF documents)
- .xlsx files (Excel spreadsheets)
- .pptx.pdf files (converted PowerPoint)

Output: Append to existing poc05_v02_optimized_chunks table
"""

import sys
import os
import json
from datetime import datetime

sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'helper_function'))

from databricks.connect import DatabricksSession
from pyspark.sql import Row


def install_libraries():
    """Install required libraries"""
    print("Installing advanced libraries for multimodal processing...")
    import subprocess
    libs = [
        "python-docx",
        "pythainlp",
        "openpyxl",
        "pdfplumber",
        "PyMuPDF",  # fitz - for images from PDF
        "pillow",   # Image processing
        "pandas",   # Better table handling
    ]
    for lib in libs:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", lib, "--quiet"])
        except:
            pass
    print("Libraries installed")


def extract_keywords(text, lang='en'):
    """Extract keywords from text"""
    words = text.lower().split()

    tech_keywords = {'table', 'column', 'log', 'id', 'status', 'data', 'submission',
                     'response', 'database', 'key', 'index', 'tbl_', 'bot', 'rdt'}

    thai_keywords = {'ตาราง', 'คอลัมน์', 'ข้อมูล', 'สถานะ', 'บันทึก', 'ธปท',
                     'ส่ง', 'ตอบกลับ', 'ระบบ', 'ฐานข้อมูล'}

    found = set()
    for word in words:
        if any(kw in word for kw in tech_keywords):
            found.add(word)
        if any(kw in word for kw in thai_keywords):
            found.add(word)

    return list(found)[:10]


def enhance_chunk_bilingual(chunk_text, section, doc_name, chunk_type='text'):
    """Add bilingual metadata to chunks"""
    en_keywords = extract_keywords(chunk_text, 'en')
    th_keywords = extract_keywords(chunk_text, 'th')

    enhanced = f"""[SECTION] {section}
KEYWORDS: {', '.join(en_keywords)}
คำสำคัญ: {', '.join(th_keywords)}
DOCUMENT: {doc_name}

{chunk_text}"""

    return enhanced


def extract_pdf_advanced(pdf_bytes, filename):
    """
    Advanced PDF extraction with:
    - Text extraction (pdfplumber)
    - Table detection and extraction
    - Image/diagram detection and description
    """
    try:
        import pdfplumber
        import fitz  # PyMuPDF
        from io import BytesIO
        import pandas as pd

        chunks = []
        current_section = "Introduction"

        # Open with both libraries for comprehensive extraction
        pdf_pdfplumber = pdfplumber.open(BytesIO(pdf_bytes))
        pdf_fitz = fitz.open(stream=pdf_bytes, filetype="pdf")

        all_text = []
        image_count = 0

        for page_num in range(len(pdf_pdfplumber.pages)):
            try:
                # pdfplumber for text and tables
                page_plumber = pdf_pdfplumber.pages[page_num]
                page_fitz = pdf_fitz[page_num]

                # Extract text
                text = page_plumber.extract_text()
                if text:
                    all_text.append(f"[PAGE {page_num + 1}]\n{text}")

                # Extract tables with better formatting
                tables = page_plumber.extract_tables()
                if tables:
                    for i, table in enumerate(tables, 1):
                        if table and len(table) > 1:
                            # Convert to pandas for better handling
                            try:
                                df = pd.DataFrame(table[1:], columns=table[0])
                                df = df.fillna('')

                                # Create table text
                                table_text = f"[TABLE] Page {page_num + 1}, Table {i}\n"
                                table_text += f"COLUMNS: {', '.join([str(c) for c in df.columns])}\n\n"
                                table_text += df.to_string(index=False, max_rows=50)

                                enhanced = enhance_chunk_bilingual(table_text, f"Page {page_num + 1}", filename, 'table')
                                chunks.append({
                                    'content': enhanced,
                                    'content_type': 'table',
                                    'section': f"Page {page_num + 1}",
                                    'char_count': len(enhanced),
                                    'original_char_count': len(table_text),
                                    'table_columns': ', '.join([str(c) for c in df.columns])
                                })
                            except Exception as e:
                                # Fallback to simple table formatting
                                headers = table[0] if table else []
                                table_text = f"[TABLE] Page {page_num + 1}, Table {i}\n"
                                table_text += " | ".join([str(h) for h in headers if h]) + "\n"

                                for row in table[1:]:
                                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                    table_text += row_text + "\n"

                                enhanced = enhance_chunk_bilingual(table_text, f"Page {page_num + 1}", filename, 'table')
                                chunks.append({
                                    'content': enhanced,
                                    'content_type': 'table',
                                    'section': f"Page {page_num + 1}",
                                    'char_count': len(enhanced),
                                    'original_char_count': len(table_text)
                                })

                # Extract images/diagrams (metadata only, not actual image data)
                images = page_fitz.get_images()
                if images:
                    for img_index, img in enumerate(images):
                        image_count += 1
                        try:
                            xref = img[0]
                            img_info = pdf_fitz.extract_image(xref)

                            # Create image reference chunk
                            img_text = f"[IMAGE/DIAGRAM] Page {page_num + 1}, Image {img_index + 1}\n"
                            img_text += f"Type: {img_info.get('ext', 'unknown')}\n"
                            img_text += f"Size: {img_info.get('width', 0)}x{img_info.get('height', 0)}\n"
                            img_text += f"Context: Located on page {page_num + 1}\n"
                            img_text += f"Note: This may be a diagram, flowchart, or illustration\n"

                            # Try to get surrounding text for context
                            if text:
                                context_words = text[:200] if len(text) > 200 else text
                                img_text += f"Surrounding text: {context_words}\n"

                            enhanced = enhance_chunk_bilingual(img_text, f"Page {page_num + 1}", filename, 'image')
                            chunks.append({
                                'content': enhanced,
                                'content_type': 'image',
                                'section': f"Page {page_num + 1}",
                                'char_count': len(enhanced),
                                'original_char_count': len(img_text)
                            })
                        except Exception as e:
                            continue

            except Exception as e:
                print(f"      Error extracting page {page_num + 1}: {e}")
                continue

        # Chunk the text
        full_text = "\n".join(all_text)
        if full_text:
            # Split into 800 char chunks with 200 char overlap
            text_chunks = []
            start = 0
            while start < len(full_text):
                end = start + 800
                chunk_text = full_text[start:end]

                enhanced = enhance_chunk_bilingual(chunk_text, current_section, filename, 'text')
                text_chunks.append({
                    'content': enhanced,
                    'content_type': 'text',
                    'section': current_section,
                    'char_count': len(enhanced),
                    'original_char_count': len(chunk_text)
                })

                start = end - 200  # 200 char overlap

            chunks.extend(text_chunks)

        pdf_pdfplumber.close()
        pdf_fitz.close()

        return chunks

    except Exception as e:
        print(f"    ERROR processing PDF {filename}: {e}")
        import traceback
        traceback.print_exc()
        return []


def extract_xlsx_advanced(xlsx_bytes, filename):
    """
    Advanced Excel extraction with:
    - Multiple sheets handling
    - Table detection with proper headers
    - Image/chart detection (metadata)
    - Better data type handling
    """
    try:
        from openpyxl import load_workbook
        from io import BytesIO
        import pandas as pd

        chunks = []
        wb = load_workbook(BytesIO(xlsx_bytes), data_only=True)

        print(f"      Found {len(wb.sheetnames)} sheets: {', '.join(wb.sheetnames)}")

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]

            # Get all rows
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue

            # Try to identify header row (first non-empty row)
            headers = None
            data_start = 0

            for i, row in enumerate(rows):
                if any(cell for cell in row):
                    headers = [str(cell) if cell else f"Col{j}" for j, cell in enumerate(row)]
                    data_start = i + 1
                    break

            if not headers:
                continue

            # Convert to pandas for better handling
            try:
                data_rows = []
                for row in rows[data_start:]:
                    if any(cell for cell in row):
                        data_rows.append(row)

                if data_rows:
                    df = pd.DataFrame(data_rows, columns=headers)
                    df = df.fillna('')

                    # Chunk every 50 rows
                    chunk_size = 50
                    for chunk_idx in range(0, len(df), chunk_size):
                        chunk_df = df.iloc[chunk_idx:chunk_idx + chunk_size]

                        table_text = f"[TABLE] Sheet: {sheet_name}"
                        if chunk_idx > 0:
                            table_text += f" (Rows {chunk_idx + 1}-{chunk_idx + len(chunk_df)})"
                        table_text += f"\nCOLUMNS: {', '.join([str(c) for c in df.columns])}\n\n"
                        table_text += chunk_df.to_string(index=False, max_rows=50)

                        enhanced = enhance_chunk_bilingual(table_text, sheet_name, filename, 'table')
                        chunks.append({
                            'content': enhanced,
                            'content_type': 'table',
                            'section': sheet_name,
                            'char_count': len(enhanced),
                            'original_char_count': len(table_text),
                            'table_columns': ', '.join([str(c) for c in df.columns])
                        })

            except Exception as e:
                # Fallback to simple processing
                table_text = f"[TABLE] Sheet: {sheet_name}\nCOLUMNS: {', '.join(headers)}\n"
                table_text += " | ".join(headers) + "\n"

                row_count = 0
                for row in rows[data_start:]:
                    if any(cell for cell in row):
                        row_text = " | ".join([str(cell) if cell else "" for cell in row])
                        table_text += row_text + "\n"
                        row_count += 1

                        # Chunk every 50 rows
                        if row_count >= 50:
                            enhanced = enhance_chunk_bilingual(table_text, sheet_name, filename, 'table')
                            chunks.append({
                                'content': enhanced,
                                'content_type': 'table',
                                'section': sheet_name,
                                'char_count': len(enhanced),
                                'original_char_count': len(table_text),
                                'table_columns': ', '.join(headers)
                            })

                            # Start new chunk
                            table_text = f"[TABLE] Sheet: {sheet_name} (continued)\nCOLUMNS: {', '.join(headers)}\n"
                            table_text += " | ".join(headers) + "\n"
                            row_count = 0

                # Add remaining rows
                if row_count > 0:
                    enhanced = enhance_chunk_bilingual(table_text, sheet_name, filename, 'table')
                    chunks.append({
                        'content': enhanced,
                        'content_type': 'table',
                        'section': sheet_name,
                        'char_count': len(enhanced),
                        'original_char_count': len(table_text),
                        'table_columns': ', '.join(headers)
                    })

            # Check for images/charts in sheet
            if hasattr(sheet, '_images') and sheet._images:
                for img_idx, img in enumerate(sheet._images, 1):
                    img_text = f"[IMAGE/CHART] Sheet: {sheet_name}, Image {img_idx}\n"
                    img_text += f"Type: Embedded image or chart\n"
                    img_text += f"Note: This sheet contains visual elements (charts, diagrams, or images)\n"

                    enhanced = enhance_chunk_bilingual(img_text, sheet_name, filename, 'image')
                    chunks.append({
                        'content': enhanced,
                        'content_type': 'image',
                        'section': sheet_name,
                        'char_count': len(enhanced),
                        'original_char_count': len(img_text)
                    })

        return chunks

    except Exception as e:
        print(f"    ERROR processing Excel {filename}: {e}")
        import traceback
        traceback.print_exc()
        return []


def extract_docx(doc_bytes, filename):
    """Extract from DOCX (reuse from 02_process script)"""
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(doc_bytes))
        chunks = []
        current_section = "Introduction"
        current_text = []
        current_size = 0
        table_count = 0

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect section headers
            if text.startswith('#') or (para.style.name.startswith('Heading') and len(text) < 100):
                # Save previous section
                if current_text:
                    chunk_content = '\n'.join(current_text)
                    enhanced = enhance_chunk_bilingual(chunk_content, current_section, filename, 'text')
                    chunks.append({
                        'content': enhanced,
                        'content_type': 'text',
                        'section': current_section,
                        'char_count': len(enhanced),
                        'original_char_count': len(chunk_content)
                    })

                # Start new section
                current_section = text.replace('#', '').strip()
                current_text = []
                current_size = 0
                continue

            # Add text to current chunk
            current_text.append(text)
            current_size += len(text)

            # Create chunk if size exceeded
            if current_size > 800:
                chunk_content = '\n'.join(current_text)
                enhanced = enhance_chunk_bilingual(chunk_content, current_section, filename, 'text')
                chunks.append({
                    'content': enhanced,
                    'content_type': 'text',
                    'section': current_section,
                    'char_count': len(enhanced),
                    'original_char_count': len(chunk_content)
                })

                # OVERLAP: Keep last 2 sentences (~200 chars)
                overlap_sentences = current_text[-2:] if len(current_text) > 2 else current_text[-1:]
                current_text = overlap_sentences
                current_size = sum(len(s) for s in current_text)

        # Final text chunk
        if current_text:
            chunk_content = '\n'.join(current_text)
            enhanced = enhance_chunk_bilingual(chunk_content, current_section, filename, 'text')
            chunks.append({
                'content': enhanced,
                'content_type': 'text',
                'section': current_section,
                'char_count': len(enhanced),
                'original_char_count': len(chunk_content)
            })

        # Process tables
        for table in doc.tables:
            table_count += 1
            try:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                table_data = []

                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(' | '.join(row_data))

                if headers and table_data:
                    column_list = ', '.join(headers)
                    table_text = f"""TABLE {table_count}: {current_section}
COLUMNS: {column_list}
Headers: {' | '.join(headers)}
{chr(10).join(table_data[:20])}"""

                    enhanced = enhance_chunk_bilingual(table_text, current_section, filename, 'table')

                    chunks.append({
                        'content': enhanced,
                        'content_type': 'table',
                        'section': current_section,
                        'char_count': len(enhanced),
                        'original_char_count': len(table_text),
                        'table_columns': column_list
                    })

            except Exception as e:
                print(f"      Table {table_count} extraction error: {e}")
                continue

        return chunks

    except Exception as e:
        print(f"    ERROR processing DOCX {filename}: {e}")
        return []


def main():
    print("=" * 80)
    print("POC05_V02 - PROCESS ALL DOCUMENT FORMATS")
    print("=" * 80)

    install_libraries()

    # Load config
    config_path = os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'databricks_config', 'config.json')
    with open(config_path, 'r') as f:
        config = json.load(f)

    # Use lineage cluster for better performance
    cluster_id = config['databricks'].get('lineage_cluster_id', config['databricks']['cluster_id'])

    print(f"\n[1] Connecting to Databricks (cluster: {cluster_id})...")
    spark = DatabricksSession.builder.remote(
        host=config['databricks']['host'],
        token=config['databricks']['token'],
        cluster_id=cluster_id
    ).getOrCreate()
    print("    Connected")

    volume_path = config['rag_config']['volume_path']
    print(f"\n[2] Scanning documents in {volume_path}...")

    # List documents
    files_df = spark.sql(f"LIST '{volume_path}'")
    all_files = [row['name'] for row in files_df.collect()]

    # Filter document files (all formats)
    doc_files = [f for f in all_files if f.lower().endswith(('.docx', '.pdf', '.xlsx', '.doc'))]

    # Count by type
    docx_files = [f for f in doc_files if f.lower().endswith('.docx')]
    pdf_files = [f for f in doc_files if f.lower().endswith('.pdf')]
    xlsx_files = [f for f in doc_files if f.lower().endswith('.xlsx')]

    print(f"    Found {len(doc_files)} documents:")
    print(f"      - {len(docx_files)} .docx files")
    print(f"      - {len(pdf_files)} .pdf files")
    print(f"      - {len(xlsx_files)} .xlsx files")

    print(f"\n[3] Processing all document formats...")
    print(f"    - Semantic chunking (800 char max)")
    print(f"    - 200 char overlap for context")
    print(f"    - Bilingual keywords (Thai + English)")
    print(f"    - Table extraction for all formats")

    all_chunks = []
    processed_count = 0
    text_chunks_list = []
    table_chunks_list = []

    stats = {
        'docx': {'success': 0, 'failed': 0, 'chunks': 0},
        'pdf': {'success': 0, 'failed': 0, 'chunks': 0},
        'xlsx': {'success': 0, 'failed': 0, 'chunks': 0}
    }

    for filename in doc_files:
        try:
            file_path = f"{volume_path}/{filename}"

            # Read file
            file_df = spark.read.format("binaryFile").load(file_path)
            file_bytes = file_df.collect()[0]['content']

            # Process based on file type
            chunks = []
            file_type = None

            if filename.lower().endswith('.docx'):
                chunks = extract_docx(file_bytes, filename)
                file_type = 'docx'
            elif filename.lower().endswith('.pdf'):
                chunks = extract_pdf_advanced(file_bytes, filename)
                file_type = 'pdf'
            elif filename.lower().endswith('.xlsx'):
                chunks = extract_xlsx_advanced(file_bytes, filename)
                file_type = 'xlsx'

            if chunks and file_type:
                for i, chunk in enumerate(chunks):
                    chunk_id = f"poc05_v02_all_{processed_count}_{i}"
                    chunk_row = {
                        'id': chunk_id,
                        'content': chunk['content'],
                        'content_type': chunk['content_type'],
                        'section': chunk['section'],
                        'char_count': chunk['char_count'],
                        'original_char_count': chunk['original_char_count'],
                        'source_document': filename,
                        'table_columns': chunk.get('table_columns', None),
                        'processing_method': f'semantic_chunking_overlap_bilingual_{file_type}'
                    }
                    all_chunks.append(chunk_row)

                    if chunk['content_type'] == 'text':
                        text_chunks_list.append(chunk)
                    else:
                        table_chunks_list.append(chunk)

                stats[file_type]['success'] += 1
                stats[file_type]['chunks'] += len(chunks)
                processed_count += 1
                print(f"  [{processed_count}/{len(doc_files)}] {filename}")
                print(f"    Extracted {len(chunks)} chunks ({file_type.upper()})")
            else:
                if file_type:
                    stats[file_type]['failed'] += 1
                print(f"  [{processed_count}/{len(doc_files)}] {filename}")
                print(f"    No chunks extracted")

        except Exception as e:
            file_type = 'docx' if filename.endswith('.docx') else 'pdf' if filename.endswith('.pdf') else 'xlsx'
            if file_type in stats:
                stats[file_type]['failed'] += 1
            print(f"  ERROR processing {filename}: {e}")
            continue

    # Save to table
    print(f"\n[4] Processing complete:")
    print(f"    Documents processed: {processed_count}/{len(doc_files)}")
    print(f"    Total chunks: {len(all_chunks)}")
    print(f"\n    By format:")
    print(f"      DOCX: {stats['docx']['success']} files, {stats['docx']['chunks']} chunks ({stats['docx']['failed']} failed)")
    print(f"      PDF:  {stats['pdf']['success']} files, {stats['pdf']['chunks']} chunks ({stats['pdf']['failed']} failed)")
    print(f"      XLSX: {stats['xlsx']['success']} files, {stats['xlsx']['chunks']} chunks ({stats['xlsx']['failed']} failed)")

    if text_chunks_list:
        avg_text_chars = sum(c['original_char_count'] for c in text_chunks_list) // len(text_chunks_list)
        print(f"\n    Text chunks: {len(text_chunks_list)} (avg {avg_text_chars} chars)")

    if table_chunks_list:
        avg_table_chars = sum(c['original_char_count'] for c in table_chunks_list) // len(table_chunks_list)
        print(f"    Table chunks: {len(table_chunks_list)} (avg {avg_table_chars} chars)")

    # Create DataFrame and save
    output_table = "sandbox.rdt_knowledge.poc05_v02_all_formats_chunks"
    print(f"\n[5] Saving to table: {output_table}")
    chunks_df = spark.createDataFrame([Row(**chunk) for chunk in all_chunks])
    chunks_df.write.format("delta").mode("overwrite").saveAsTable(output_table)
    print(f"    Saved {len(all_chunks)} chunks")

    # Update config
    config['rag_config']['poc05_v02_all_formats_table'] = output_table
    config['rag_config']['poc05_v02_all_formats_count'] = len(all_chunks)
    config['rag_config']['poc05_v02_all_formats_stats'] = stats
    config['rag_config']['poc05_v02_all_formats_created_at'] = datetime.now().isoformat()

    with open(config_path, 'w') as f:
        json.dump(config, f, indent=2)

    print("\n" + "=" * 80)
    print("ALL FORMATS PROCESSING COMPLETE")
    print("=" * 80)
    print(f"\nTable: {output_table}")
    print(f"Chunks: {len(all_chunks)}")
    print(f"Documents: {processed_count}/{len(doc_files)}")
    print(f"\nFormat breakdown:")
    print(f"  DOCX: {stats['docx']['chunks']} chunks from {stats['docx']['success']} files")
    print(f"  PDF:  {stats['pdf']['chunks']} chunks from {stats['pdf']['success']} files")
    print(f"  XLSX: {stats['xlsx']['chunks']} chunks from {stats['xlsx']['success']} files")
    print(f"\nNext: Add metadata and create vector index for all formats")

    spark.stop()


if __name__ == "__main__":
    main()

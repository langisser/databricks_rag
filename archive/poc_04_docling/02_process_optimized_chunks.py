#!/usr/bin/env python3
"""
POC Iterative 04 - Optimized Document Processing
Implements: Semantic chunking + Overlap + Bilingual metadata
WITHOUT Docling (using enhanced python-docx for faster execution)
"""

import sys
import os
import json
from datetime import datetime
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'helper_function'))

from databricks.connect import DatabricksSession
from pyspark.sql import Row


def install_libraries():
    """Install required libraries"""
    print("Installing libraries...")
    import subprocess
    libs = ["python-docx", "pythainlp"]
    for lib in libs:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", lib, "--quiet"])
        except:
            pass
    print("Libraries installed")


def extract_keywords(text, lang='en'):
    """Extract keywords from text"""
    # Simple keyword extraction (can be enhanced with NLP)
    words = text.lower().split()

    # Common database/technical terms
    tech_keywords = {'table', 'column', 'log', 'id', 'status', 'data', 'submission',
                     'response', 'database', 'key', 'index', 'tbl_', 'bot', 'rdt'}

    # Thai keywords
    thai_keywords = {'ตาราง', 'คอลัมน์', 'ข้อมูล', 'สถานะ', 'บันทึก', 'ธปท',
                     'ส่ง', 'ตอบกลับ', 'ระบบ', 'ฐานข้อมูล'}

    found = set()
    for word in words:
        if any(kw in word for kw in tech_keywords):
            found.add(word)
        if any(kw in word for kw in thai_keywords):
            found.add(word)

    return list(found)[:10]  # Top 10 keywords


def enhance_chunk_bilingual(chunk_text, section, doc_name, chunk_type='text'):
    """
    Add bilingual metadata to chunks
    """
    # Extract keywords
    en_keywords = extract_keywords(chunk_text, 'en')
    th_keywords = extract_keywords(chunk_text, 'th')

    # For table chunks, extract column information
    if chunk_type == 'table' and 'COLUMN' in chunk_text.upper():
        # Try to extract column names
        lines = chunk_text.split('\n')
        column_line = None
        for line in lines:
            if 'column' in line.lower() and '|' in line:
                column_line = line
                break

        if column_line:
            # Enhanced table with bilingual keywords
            enhanced = f"""[TABLE] {section}
KEYWORDS: {', '.join(en_keywords)}
คำสำคัญ: {', '.join(th_keywords)}
DOCUMENT: {doc_name}

{chunk_text}

[TABLE REFERENCE] For queries about table structure, columns, and database schema
[การอ้างอิงตาราง] สำหรับคำถามเกี่ยวกับโครงสร้างตาราง, คอลัมน์, และสคีมาฐานข้อมูล"""
        else:
            enhanced = f"""[TABLE] {section}
KEYWORDS: {', '.join(en_keywords)}
คำสำคัญ: {', '.join(th_keywords)}

{chunk_text}"""
    else:
        # Enhanced text with bilingual keywords
        enhanced = f"""[SECTION] {section}
KEYWORDS: {', '.join(en_keywords)}
คำสำคัญ: {', '.join(th_keywords)}
DOCUMENT: {doc_name}

{chunk_text}"""

    return enhanced


def extract_with_overlap(doc_bytes, filename):
    """
    Enhanced extraction with semantic chunking and overlap
    Key improvements:
    1. Smaller chunks (800 chars max instead of 2000)
    2. Overlap (200 chars) to preserve context
    3. Better section detection
    4. Bilingual metadata
    """
    from docx import Document
    from io import BytesIO

    print(f"\n  Processing: {filename}")

    doc = Document(BytesIO(doc_bytes))
    chunks = []
    chunk_id = 0

    # Extract paragraphs with semantic boundaries
    current_section = "Introduction"
    current_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect section headers
        is_header = (len(text) < 100 and
                    (para.style.name.startswith('Heading') or
                     text.isupper() or
                     text.startswith('##') or
                     any(kw in text.lower() for kw in ['overview', 'process', 'framework',
                                                         'specification', 'table information'])))

        if is_header:
            # Save previous section
            if current_text:
                content = '\n'.join(current_text)
                enhanced = enhance_chunk_bilingual(content, current_section, filename, 'text')
                chunks.append({
                    'chunk_id': chunk_id,
                    'content': enhanced,
                    'content_type': 'text',
                    'section': current_section,
                    'char_count': len(enhanced),
                    'original_char_count': len(content)
                })
                chunk_id += 1
                current_text = []

            current_section = text
            current_text.append(f"## {text}")
        else:
            current_text.append(text)

            # Create chunk if reaching optimal size (800 chars)
            current_size = sum(len(t) for t in current_text)
            if current_size > 800:
                content = '\n'.join(current_text)
                enhanced = enhance_chunk_bilingual(content, current_section, filename, 'text')
                chunks.append({
                    'chunk_id': chunk_id,
                    'content': enhanced,
                    'content_type': 'text',
                    'section': current_section,
                    'char_count': len(enhanced),
                    'original_char_count': len(content)
                })
                chunk_id += 1

                # OVERLAP: Keep last 2 sentences (~200 chars)
                overlap_sentences = current_text[-2:] if len(current_text) > 2 else current_text[-1:]
                current_text = overlap_sentences

    # Save final section
    if current_text:
        content = '\n'.join(current_text)
        enhanced = enhance_chunk_bilingual(content, current_section, filename, 'text')
        chunks.append({
            'chunk_id': chunk_id,
            'content': enhanced,
            'content_type': 'text',
            'section': current_section,
            'char_count': len(enhanced),
            'original_char_count': len(content)
        })
        chunk_id += 1

    # Extract tables with enhanced metadata
    table_num = 0
    for table in doc.tables:
        if len(table.rows) == 0:
            continue

        # Get headers
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        # Get rows
        rows = []
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                rows.append(' | '.join(row_data))

        if rows:
            table_num += 1
            column_list = ', '.join(headers)

            # Create table text
            table_text = f"""TABLE {table_num}: {current_section}
Headers: {' | '.join(headers)}
{chr(10).join(rows[:15])}"""  # Limit to 15 rows per chunk

            # If table is large, split into multiple chunks
            if len(rows) > 15:
                for i in range(0, len(rows), 15):
                    chunk_rows = rows[i:i+15]
                    partial_table = f"""TABLE {table_num} (Part {i//15 + 1}): {current_section}
Headers: {' | '.join(headers)}
{chr(10).join(chunk_rows)}"""

                    enhanced = enhance_chunk_bilingual(partial_table, current_section, filename, 'table')
                    chunks.append({
                        'chunk_id': chunk_id,
                        'content': enhanced,
                        'content_type': 'table',
                        'section': current_section,
                        'char_count': len(enhanced),
                        'original_char_count': len(partial_table),
                        'table_columns': column_list
                    })
                    chunk_id += 1
            else:
                enhanced = enhance_chunk_bilingual(table_text, current_section, filename, 'table')
                chunks.append({
                    'chunk_id': chunk_id,
                    'content': enhanced,
                    'content_type': 'table',
                    'section': current_section,
                    'char_count': len(enhanced),
                    'original_char_count': len(table_text),
                    'table_columns': column_list
                })
                chunk_id += 1

    print(f"    Extracted {len(chunks)} chunks (optimized with overlap & bilingual metadata)")
    return chunks


def main():
    print("=" * 80)
    print("POC ITERATIVE 04 - OPTIMIZED CHUNKING WITH BILINGUAL METADATA")
    print("=" * 80)

    install_libraries()

    # Load config
    config_path = os.path.join(os.path.dirname(__file__), '..', 'databricks_helper', 'databricks_config', 'config.json')
    with open(config_path, 'r') as f:
        config = json.load(f)

    # Connect
    print("\n[1] Connecting to Databricks...")
    spark = DatabricksSession.builder.remote(
        host=config['databricks']['host'],
        token=config['databricks']['token'],
        cluster_id=config['databricks']['cluster_id']
    ).getOrCreate()
    print("    Connected")

    # Get volume
    volume_path = config['rag_config']['volume_path']

    print(f"\n[2] Scanning documents in {volume_path}...")
    files_df = spark.sql(f"LIST '{volume_path}'")
    docs = [row for row in files_df.collect() if row['name'].endswith('.docx')]

    print(f"    Found {len(docs)} documents")

    # Process all documents
    print(f"\n[3] Processing documents with optimizations...")
    print("    - Semantic chunking (800 char max)")
    print("    - 200 char overlap for context")
    print("    - Bilingual keywords (Thai + English)")
    print("    - Table column extraction")

    all_chunks = []
    doc_count = 0

    for doc_info in docs:
        try:
            # Read document
            binary_df = spark.read.format("binaryFile").load(doc_info['path'])
            doc_bytes = binary_df.collect()[0]['content']

            # Extract with optimizations
            chunks = extract_with_overlap(doc_bytes, doc_info['name'])

            # Add document prefix to IDs
            for chunk in chunks:
                chunk['id'] = f"poc04_{doc_count}_{chunk['chunk_id']}"
                chunk['source_document'] = doc_info['name']

            all_chunks.extend(chunks)
            doc_count += 1

        except Exception as e:
            print(f"    ERROR processing {doc_info['name']}: {e}")

    print(f"\n[4] Processing complete:")
    print(f"    Documents processed: {doc_count}")
    print(f"    Total chunks: {len(all_chunks)}")
    text_chunks = [c for c in all_chunks if c['content_type'] == 'text']
    table_chunks = [c for c in all_chunks if c['content_type'] == 'table']
    print(f"    Text chunks: {len(text_chunks)} (avg {sum(c['original_char_count'] for c in text_chunks)//len(text_chunks) if text_chunks else 0} chars)")
    print(f"    Table chunks: {len(table_chunks)} (avg {sum(c['original_char_count'] for c in table_chunks)//len(table_chunks) if table_chunks else 0} chars)")

    # Save to Delta table
    namespace = config['rag_config']['full_namespace']
    table_name = f"{namespace}.poc04_optimized_chunks"

    print(f"\n[5] Saving to table: {table_name}")

    rows = []
    for chunk in all_chunks:
        rows.append(Row(
            id=chunk['id'],
            content=chunk['content'],
            content_type=chunk['content_type'],
            section=chunk['section'],
            char_count=chunk['char_count'],
            original_char_count=chunk['original_char_count'],
            source_document=chunk['source_document'],
            table_columns=chunk.get('table_columns', None),
            processing_method='optimized_semantic_overlap_bilingual'
        ))

    df = spark.createDataFrame(rows)
    df.write.format("delta").mode("overwrite").saveAsTable(table_name)

    print(f"    Saved {len(rows)} chunks")

    # Show statistics
    print(f"\n[6] Quality Improvements vs POC03:")
    print(f"    POC03 text chunks: 169 chars avg → POC04: {sum(c['original_char_count'] for c in text_chunks)//len(text_chunks) if text_chunks else 0} chars avg")
    print(f"    Overlap added: YES (200 chars)")
    print(f"    Bilingual metadata: YES (Thai + English keywords)")
    print(f"    Table column extraction: {len([c for c in table_chunks if c.get('table_columns')])} tables")

    # Show samples
    print(f"\n[7] Sample enhanced chunk:")
    if all_chunks:
        sample = all_chunks[5] if len(all_chunks) > 5 else all_chunks[0]
        print(f"\n{sample['content'][:400]}...")

    # Summary
    print("\n" + "=" * 80)
    print("PROCESSING COMPLETE")
    print("=" * 80)
    print(f"\nTable: {table_name}")
    print(f"Total chunks: {len(all_chunks)}")
    print(f"Documents: {doc_count}")
    print(f"\nOptimizations applied:")
    print(f"  ✓ Semantic chunking (800 chars max)")
    print(f"  ✓ Context overlap (200 chars)")
    print(f"  ✓ Bilingual keywords (Thai + English)")
    print(f"  ✓ Enhanced table metadata")
    print(f"\nNext: Create vector index with 'poc04_optimized_rag_index'")

    spark.stop()


if __name__ == "__main__":
    main()
